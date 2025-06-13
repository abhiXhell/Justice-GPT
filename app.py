import streamlit as st
import os
from src.rag_engine import RAGEngine
import tempfile
from pypdf import PdfReader
import json
from datetime import datetime
from docx import Document

CHAT_HISTORY_PATH = "data/chat_history.json"
TRAINED_MODEL_PATH = "data/trained_model"
OUTPUTS_DIR = "outputs"

# Helper functions for chat history

def load_chat_history():
    if os.path.exists(CHAT_HISTORY_PATH):
        with open(CHAT_HISTORY_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    return []

def save_chat_history(history):
    with open(CHAT_HISTORY_PATH, "w", encoding="utf-8") as f:
        json.dump(history, f, ensure_ascii=False, indent=2)

def export_response_to_docx(response, user_question):
    if not os.path.exists(OUTPUTS_DIR):
        os.makedirs(OUTPUTS_DIR)
    doc = Document()
    doc.add_heading('Complaint Summary', 0)
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"User Question: {user_question}")
    # Try to split response into sections
    incident, legal_basis, suggested_actions = '', '', ''
    lower = response.lower()
    if 'incident:' in lower and 'legal basis:' in lower and 'suggested actions:' in lower:
        try:
            incident = response.split('Incident:')[1].split('Legal Basis:')[0].strip()
            legal_basis = response.split('Legal Basis:')[1].split('Suggested Actions:')[0].strip()
            suggested_actions = response.split('Suggested Actions:')[1].strip()
        except Exception:
            incident, legal_basis, suggested_actions = '', '', ''
    if incident:
        doc.add_heading('Incident', level=1)
        doc.add_paragraph(incident)
    if legal_basis:
        doc.add_heading('Legal Basis', level=1)
        doc.add_paragraph(legal_basis)
    if suggested_actions:
        doc.add_heading('Suggested Actions', level=1)
        doc.add_paragraph(suggested_actions)
    if not (incident or legal_basis or suggested_actions):
        doc.add_paragraph(response)
    filename = os.path.join(OUTPUTS_DIR, f"complaint_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    doc.save(filename)
    return filename

# Initialize session state
if 'rag_engine' not in st.session_state:
    st.session_state.rag_engine = RAGEngine()
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = load_chat_history()
if 'model_trained' not in st.session_state:
    st.session_state.model_trained = os.path.exists(TRAINED_MODEL_PATH)

# Page config
st.set_page_config(
    page_title="Indian Labour & Consumer Law AI Assistant",
    page_icon="⚖️",
    layout="wide"
)

# Title and description
st.title("⚖️ Indian Labour & Consumer Law AI Assistant")
st.markdown("""
This AI assistant specializes in Indian Labour and Consumer Court law. 
Upload relevant legal documents and ask questions to get accurate, domain-specific answers.
""")

# Sidebar for document upload and further training
with st.sidebar:
    st.header("Document Upload")
    uploaded_files = st.file_uploader(
        "Upload additional legal documents (PDF)",
        type=['pdf'],
        accept_multiple_files=True
    )
    if uploaded_files:
        for uploaded_file in uploaded_files:
            # Save uploaded file temporarily
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                tmp_path = tmp_file.name
            # Process the document
            st.session_state.rag_engine.process_document(tmp_path)
            # Clean up temporary file
            os.unlink(tmp_path)
            st.success(f'Processed {uploaded_file.name}')
    if st.button("Clear All Documents"):
        st.session_state.rag_engine.clear_documents()
        st.session_state.chat_history = []
        save_chat_history([])
        st.success("All documents cleared!")

    st.markdown("---")
    st.header("Add More Training Documents")
    train_files = st.file_uploader(
        "Upload PDFs to further train the model (optional)",
        type=['pdf'],
        accept_multiple_files=True,
        key="train_files"
    )
    if train_files:
        for train_file in train_files:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                tmp_file.write(train_file.getvalue())
                tmp_path = tmp_file.name
            with st.spinner(f"Training model on {train_file.name}..."):
                reader = PdfReader(tmp_path)
                text = ""
                for page in reader.pages:
                    text += page.extract_text()
                st.session_state.rag_engine.llm_interface.train_model(text)
            os.unlink(tmp_path)
            st.success(f"Model further trained on {train_file.name}")

# Main chat interface (always available if model is trained)
if st.session_state.model_trained:
    st.header("Chat Interface")
    # Display chat history
    for message in st.session_state.chat_history:
        with st.chat_message(message["role"]):
            st.write(message["content"])
    # Chat input
    if prompt := st.chat_input("Ask a question about Indian Labour or Consumer Court law"):
        st.session_state.chat_history.append({"role": "user", "content": prompt})
        save_chat_history(st.session_state.chat_history)
        with st.chat_message("user"):
            st.write(prompt)
        # Generate and display assistant response (no spinner)
        with st.chat_message("assistant"):
            response = st.session_state.rag_engine.generate_answer(prompt)
            st.write(response)
            st.session_state.chat_history.append({"role": "assistant", "content": response})
            save_chat_history(st.session_state.chat_history)
    # Export last assistant response
    if st.session_state.chat_history and st.session_state.chat_history[-1]["role"] == "assistant":
        if st.button("Export Last Response as .docx"):
            user_q = ""
            for msg in reversed(st.session_state.chat_history):
                if msg["role"] == "user":
                    user_q = msg["content"]
                    break
            filename = export_response_to_docx(st.session_state.chat_history[-1]["content"], user_q)
            st.success(f"Exported to {filename}")
    st.markdown("---")
    st.markdown("""
    <div style='text-align: center'>
        <p>This AI assistant is designed to provide information about Indian Labour and Consumer Court law only.</p>
        <p>Please note that this is not legal advice and should not be used as a substitute for professional legal consultation.</p>
    </div>
    """, unsafe_allow_html=True)
else:
    st.info("No trained model found. Please train the model at least once using a legal document PDF.") 